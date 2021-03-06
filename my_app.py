from docx import Document
from docx.shared import Inches


document = Document()

document.add_picture(
    'me.jpg',
    width=Inches(2.0))

name = input('What is your name? ')
phone = input('What is your phone number? ')
email = input('What is your email address? ')

document.add_paragraph(
    name + ' | ' + phone + ' | ' + email)


#about me
document.add_heading('About Me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

#work experiences
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True


experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)


#more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True


        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break


#skills
document.add_heading('Skills')
skill = input('What are your skills? ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input(
        'Do you have more skills? ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter your skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break


document.save('cv.docx')